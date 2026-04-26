"""CV adaptation + Motivationsschreiben generator service.

Given a master CV (PDF/DOCX) and a job offer, produces:
  - An adapted CV (.docx) with a tailored profile section for the job
  - A Motivationsschreiben (.docx) personalised to the job

Uses python-docx for Word generation, pypdf + python-docx for parsing.
No external LLM calls — uses German templates with job-specific variables.
"""
import logging
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm
from pypdf import PdfReader

logger = logging.getLogger(__name__)

# === Storage layout ===
DATA_DIR = Path("data")
UPLOADS_DIR = DATA_DIR / "uploads"
GENERATED_DIR = DATA_DIR / "generated"
UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
GENERATED_DIR.mkdir(parents=True, exist_ok=True)


def extract_text_from_file(file_path: str, file_type: str) -> str:
    """Extract plain text from a CV file (PDF or DOCX)."""
    file_type = file_type.lower().lstrip(".")
    try:
        if file_type == "pdf":
            reader = PdfReader(file_path)
            return "\n".join((p.extract_text() or "") for p in reader.pages).strip()

        if file_type in ("docx", "doc"):
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n".join(paragraphs).strip()
    except Exception as e:
        logger.error(f"Text extraction failed for {file_path}: {e}")
    return ""


def _add_heading(doc: Document, text: str, size: int = 14, color=(232, 93, 61)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    return p


def _add_paragraph(doc: Document, text: str, size: int = 11, bold: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p


def generate_adapted_cv(
    cv_text: str,
    job_title: str,
    company_name: str,
    job_location: str,
    job_domain: str,
    output_path: str,
) -> str:
    """Generate an adapted CV Word file based on master CV text + job context."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LEBENSLAUF")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(232, 93, 61)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(f"Angepasst für: {job_title} – {company_name}")
    sub_run.italic = True
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_paragraph()  # spacer

    # Job-tailored profile section
    _add_heading(doc, "ZIEL / PROFIL")
    _add_paragraph(
        doc,
        f"Engagierte/r Bewerber/in mit Interesse an der Position als "
        f"{job_title} bei {company_name} in {job_location}. "
        f"Bringt nachweisbare Kompetenzen im Bereich {job_domain} sowie eine "
        f"hohe Motivation für anspruchsvolle Aufgaben mit.",
    )

    doc.add_paragraph()

    # Insert master CV text
    _add_heading(doc, "BERUFLICHE ERFAHRUNG & QUALIFIKATIONEN")

    cv_text = (cv_text or "").strip()
    if cv_text:
        # Split into paragraphs of reasonable length
        for line in cv_text.split("\n"):
            line = line.strip()
            if line:
                _add_paragraph(doc, line, size=10)
    else:
        _add_paragraph(
            doc,
            "[Dieser Bereich wurde aus dem hochgeladenen CV übernommen. "
            "Bitte stelle sicher, dass deine Master-CV strukturiert ist.]",
            size=10,
        )

    # Save
    doc.save(output_path)
    return output_path


def generate_motivationsschreiben(
    job_title: str,
    company_name: str,
    job_location: str,
    job_domain: str,
    job_description: str,
    output_path: str,
) -> str:
    """Generate a German Motivationsschreiben Word file."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Date (right aligned)
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_p.add_run(datetime.now().strftime("%d.%m.%Y"))
    date_run.font.size = Pt(11)

    # Recipient
    _add_paragraph(doc, company_name, bold=True)
    _add_paragraph(doc, "Personalabteilung")
    _add_paragraph(doc, job_location)

    doc.add_paragraph()

    # Subject
    subject = doc.add_paragraph()
    subj_run = subject.add_run(
        f"Bewerbung als {job_title}"
    )
    subj_run.bold = True
    subj_run.font.size = Pt(12)

    doc.add_paragraph()

    # Greeting
    _add_paragraph(doc, "Sehr geehrte Damen und Herren,")
    doc.add_paragraph()

    # Body — opening
    snippet = (job_description or "").strip()[:200]
    body_intro = (
        f"mit großem Interesse habe ich Ihre Stellenausschreibung als "
        f"{job_title} bei {company_name} gelesen. Die ausgeschriebene Position "
        f"im Bereich {job_domain} entspricht genau meinen beruflichen Zielen "
        f"und ich bewerbe mich daher mit großer Motivation um diese Stelle."
    )
    _add_paragraph(doc, body_intro)
    doc.add_paragraph()

    # Body — qualifications
    body_quals = (
        f"Durch meine bisherige Ausbildung und praktische Erfahrungen im Umfeld "
        f"{job_domain} verfüge ich über die fachlichen Kompetenzen, die für die "
        f"Position bei {company_name} relevant sind. Ich arbeite strukturiert, "
        f"lösungsorientiert und schätze sowohl Teamarbeit als auch eigenverantwortliches "
        f"Handeln. Neue Technologien und Herausforderungen begeistern mich, und ich "
        f"bringe die Bereitschaft mit, mich kontinuierlich weiterzuentwickeln."
    )
    _add_paragraph(doc, body_quals)
    doc.add_paragraph()

    # Body — why this company
    body_why = (
        f"{company_name} überzeugt mich durch das innovative Arbeitsumfeld und die "
        f"spannenden Aufgaben in {job_location}. Besonders motiviert mich die "
        f"Möglichkeit, in einem dynamischen Team einen wertvollen Beitrag zu leisten "
        f"und mein Wissen weiter auszubauen."
    )
    _add_paragraph(doc, body_why)

    if snippet:
        doc.add_paragraph()
        ref_p = doc.add_paragraph()
        ref_run = ref_p.add_run(
            f"Ihre beschriebenen Anforderungen — \"{snippet}...\" — passen "
            f"sehr gut zu meinem Profil."
        )
        ref_run.italic = True
        ref_run.font.size = Pt(10)

    doc.add_paragraph()

    # Closing
    body_close = (
        "Über die Einladung zu einem persönlichen Gespräch, in dem ich Ihnen "
        "meine Motivation und Qualifikationen näher erläutern darf, würde ich "
        "mich sehr freuen."
    )
    _add_paragraph(doc, body_close)

    doc.add_paragraph()
    _add_paragraph(doc, "Mit freundlichen Grüßen")
    doc.add_paragraph()
    _add_paragraph(doc, "[Ihr Name]", bold=True)

    doc.save(output_path)
    return output_path


def safe_filename(name: str, max_len: int = 40) -> str:
    """Make a filesystem-safe filename component."""
    cleaned = re.sub(r"[^A-Za-z0-9 _-]", "", name or "doc")
    cleaned = re.sub(r"\s+", "_", cleaned).strip("_-")
    return cleaned[:max_len] or "doc"


def generate_application_for_job(
    cv_text: str,
    job_id: int,
    job_title: str,
    company_name: str,
    job_location: str,
    job_domain: str,
    job_description: Optional[str],
) -> dict:
    """Generate both CV and Motivationsschreiben for a job. Returns paths."""
    safe_title = safe_filename(job_title)
    safe_company = safe_filename(company_name)

    cv_filename = f"CV_{safe_company}_{safe_title}_{job_id}.docx"
    mot_filename = f"Motivation_{safe_company}_{safe_title}_{job_id}.docx"

    cv_path = str(GENERATED_DIR / cv_filename)
    mot_path = str(GENERATED_DIR / mot_filename)

    generate_adapted_cv(
        cv_text=cv_text,
        job_title=job_title,
        company_name=company_name,
        job_location=job_location,
        job_domain=job_domain,
        output_path=cv_path,
    )

    generate_motivationsschreiben(
        job_title=job_title,
        company_name=company_name,
        job_location=job_location,
        job_domain=job_domain,
        job_description=job_description or "",
        output_path=mot_path,
    )

    return {"cv_path": cv_path, "motivation_path": mot_path}
